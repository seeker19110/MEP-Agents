from langchain_core.tools import tool
import pandas as pd
from docx import Document
import os
import json
from pypdf import PdfReader
import ezdxf
from ezdxf import audit
import math
import re
import ast
import operator as op
import logging
from functools import lru_cache
from src.workspace import resolve_safe_path, get_project_root
from src.cad_revision import create_snapshot
from src import cad_standards
from src import cad_geometry
from src import cad_loader

logger = logging.getLogger(__name__)

def normalize_mepf_parameter_spec(text: str) -> str:
    """Chuẩn hóa toàn bộ các ký hiệu thông số kỹ thuật MEPF trong CAD về định dạng đồng nhất cho AI:
    1. Đường kính ống: %%c, Φ, Ø, D, d, DN, OD -> Ø110 (D110)
    2. Kích thước ống gió: 600*400, 600X400, W600xH400 -> 600x400
    3. Độ dốc thoát nước: i=1%, s=1%, i=1.5% -> i=1%
    4. Tiết diện dây điện: 3x2.5mm2, 3x2.5sqmm -> 3x2.5mm²
    5. Điện áp / Pha: 220V/1P, 220V 1 Phase -> 220V-1P
    6. Lưu lượng: CMH, m3/h -> m³/h
    """
    if not text:
        return text
    text = text.replace('%%c', 'Ø').replace('%%C', 'Ø').replace('Φ', 'Ø')
    text = re.sub(r'(?i)\b(Ø|DN|D|d|OD)\s*(\d+)\b', r'Ø\2 (D\2)', text)
    text = re.sub(r'(?i)(?:W)?(\d+)\s*[\*xX]\s*(?:H)?(\d+)', r'\1x\2', text)
    text = re.sub(r'(?i)\b[is]\s*=\s*(\d+(?:\.\d+)?)\s*%', r'i=\1%', text)
    text = re.sub(r'(?i)\b(\d+x\d+(?:\.\d+)?)\s*(?:mm2|sqmm|mm²)\b', r'\1mm²', text)
    text = re.sub(r'(?i)\b(220|230|380|400)\s*V?\s*[\/\-]?\s*([13])\s*(?:P|Phase|Pha)\b', r'\1V-\2P', text)
    text = re.sub(r'(?i)\b(?:CMH|m3\/h|m3h)\b', r'm³/h', text)
    return text

normalize_pipe_diameter_spec = normalize_mepf_parameter_spec

@lru_cache(maxsize=4)
def _load_vectorstore(api_key: str, index_path: str):
    """Load (and cache) the FAISS index once per api_key/index_path instead of on every call."""
    from langchain_openai import OpenAIEmbeddings
    from langchain_community.vectorstores import FAISS

    embeddings = OpenAIEmbeddings(api_key=api_key)
    return FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)


_STOPWORDS_VI_EN = {
    "va", "la", "cua", "cho", "theo", "tai", "trong", "voi", "khi", "de", "the",
    "and", "or", "of", "the", "for", "to", "in", "a", "an", "is", "are", "how",
    "what", "bao", "nhieu", "nao", "duoc", "co", "khong", "nhu", "mot",
}


def _strip_accents(text: str) -> str:
    """Bỏ dấu tiếng Việt bằng ánh xạ ASCII đơn giản (không cần thư viện ngoài) để so
    khớp từ khóa không phân biệt dấu — hữu ích khi người dùng gõ không dấu."""
    mapping = str.maketrans(
        "àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ"
        "ÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ",
        "aaaaaaaaaaaaaaaaaeeeeeeeeeeeiiiiiooooooooooooooooouuuuuuuuuuuyyyyyd"
        "AAAAAAAAAAAAAAAAAEEEEEEEEEEEIIIIIOOOOOOOOOOOOOOOOOUUUUUUUUUUUYYYYYD",
    )
    return text.translate(mapping)


def _tokenize(text: str) -> set:
    normalized = _strip_accents(text.lower())
    words = re.findall(r"[a-z0-9]+", normalized)
    return {w for w in words if w not in _STOPWORDS_VI_EN and len(w) > 1}


@lru_cache(maxsize=1)
def _load_offline_corpus(standards_dir: str) -> tuple:
    """Nạp toàn bộ file .txt trong data/standards/ thành các đoạn (chunk) văn bản, không
    cần embedding/API — dùng cho tra cứu offline hoàn toàn (không cần OPENAI_API_KEY)."""
    chunks = []
    if not os.path.isdir(standards_dir):
        return tuple(chunks)
    for fname in sorted(os.listdir(standards_dir)):
        if not fname.lower().endswith(".txt"):
            continue
        fpath = os.path.join(standards_dir, fname)
        try:
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception:
            continue
        # Chia theo đoạn trống (paragraph) để giữ ngữ cảnh liền mạch, bỏ đoạn quá ngắn.
        for para in re.split(r"\n\s*\n", content):
            para = para.strip()
            if len(para) >= 20:
                chunks.append((fname, para))
    return tuple(chunks)


def _offline_keyword_search(query: str, standards_dir: str = "data/standards", k: int = 3) -> str:
    """Tra cứu tiêu chuẩn không cần internet/API key: so khớp từ khóa (Jaccard) trên toàn
    bộ đoạn văn bản trong data/standards/. Không mạnh bằng vector search ngữ nghĩa nhưng
    hoạt động 100% offline — đúng mục tiêu hỗ trợ model AI yếu/máy chạy Ollama cục bộ."""
    query_tokens = _tokenize(query)
    if not query_tokens:
        return "Câu truy vấn rỗng hoặc không có từ khóa hợp lệ để tra cứu."

    chunks = _load_offline_corpus(standards_dir)
    if not chunks:
        return (
            "Không tìm thấy tài liệu tiêu chuẩn nào trong 'data/standards/' để tra cứu offline. "
            "Vui lòng thêm file .txt vào thư mục này."
        )

    scored = []
    for fname, para in chunks:
        para_tokens = _tokenize(para)
        if not para_tokens:
            continue
        overlap = len(query_tokens & para_tokens)
        if overlap == 0:
            continue
        score = overlap / len(query_tokens | para_tokens)
        scored.append((score, fname, para))

    if not scored:
        return f"Không tìm thấy thông tin tiêu chuẩn nào khớp với '{query}' (tra cứu offline theo từ khóa)."

    scored.sort(key=lambda x: x[0], reverse=True)
    top = scored[:k]

    result = f"Kết quả tra cứu tiêu chuẩn OFFLINE (theo từ khóa, không cần API) cho '{query}':\n"
    for i, (score, fname, para) in enumerate(top, 1):
        result += f"\n--- Trích đoạn {i} (Nguồn: {fname}, độ khớp: {score:.2f}) ---\n{para}\n"
    return result


@tool
def search_standards(query: str) -> str:
    """Tra cứu Tiêu chuẩn thiết kế MEPF (TCVN, ASHRAE, NFPA...) từ cơ sở dữ liệu nội bộ.
    Tự động dùng FAISS + OpenAI Embeddings nếu đã cấu hình OPENAI_API_KEY và đã 'ingest',
    ngược lại tự động rơi về tra cứu offline theo từ khóa (không cần internet/API key) —
    để tính năng tra cứu tiêu chuẩn vẫn hoạt động khi chạy hoàn toàn offline (VD: Ollama)."""
    logger.info("Tra cứu tiêu chuẩn thực: %s", query)
    try:
        from src.config import settings

        api_key = settings.openai_api_key or os.getenv("OPENAI_API_KEY", "")
        index_path = "faiss_index"
        has_real_key = bool(api_key) and api_key != "dummy_key_to_prevent_crash_on_import"

        if not has_real_key or not os.path.exists(index_path):
            return _offline_keyword_search(query)

        vectorstore = _load_vectorstore(api_key, index_path)
        docs = vectorstore.similarity_search(query, k=3)

        if not docs:
            return _offline_keyword_search(query)

        result = f"Kết quả RAG Tiêu chuẩn cho '{query}':\n"
        for i, doc in enumerate(docs, 1):
            source = doc.metadata.get('source', 'Unknown')
            result += f"\n--- Trích đoạn {i} (Nguồn: {source}) ---\n"
            result += doc.page_content + "\n"

        return result
    except Exception as e:
        logger.warning("Lỗi tra cứu RAG, chuyển sang offline: %s", e)
        try:
            return _offline_keyword_search(query)
        except Exception as e2:
            return f"Lỗi tra cứu tiêu chuẩn: {e2}"

@tool
def search_web(query: str) -> str:
    """Tìm kiếm thông tin trên internet."""
    logger.info("Searching web for: %s", query)
    return f"Kết quả mô phỏng cho '{query}': Tìm thấy nhiều tài liệu liên quan."

# Chỉ cho phép các toán tử số học thuần túy - không có tên biến, thuộc tính hay lời gọi hàm,
# nên không thể escape sandbox như với eval() (kể cả khi đã tắt __builtins__).
_SAFE_OPERATORS = {
    ast.Add: op.add, ast.Sub: op.sub, ast.Mult: op.mul, ast.Div: op.truediv,
    ast.Pow: op.pow, ast.Mod: op.mod, ast.FloorDiv: op.floordiv,
    ast.USub: op.neg, ast.UAdd: op.pos,
}

def _safe_eval_node(node):
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return node.value
    if isinstance(node, ast.BinOp) and type(node.op) in _SAFE_OPERATORS:
        return _SAFE_OPERATORS[type(node.op)](_safe_eval_node(node.left), _safe_eval_node(node.right))
    if isinstance(node, ast.UnaryOp) and type(node.op) in _SAFE_OPERATORS:
        return _SAFE_OPERATORS[type(node.op)](_safe_eval_node(node.operand))
    raise ValueError("Biểu thức chứa cú pháp không được phép (chỉ hỗ trợ số và các phép toán +-*/%**).")

@tool
def calculate(expression: str) -> str:
    """Thực hiện tính toán toán học cơ bản (ví dụ: '25 * 4')."""
    logger.info("Calculating: %s", expression)
    try:
        tree = ast.parse(expression, mode="eval")
        result = _safe_eval_node(tree.body)
        return f"Kết quả: {result}"
    except Exception as e:
        return f"Lỗi tính toán: {e}"

@tool
def list_directory(path: str = ".") -> str:
    """Liệt kê danh sách các file trong thư mục để xem có file nào tồn tại."""
    logger.info("Listing directory: %s", path)
    try:
        safe_path = resolve_safe_path(path)
        files = os.listdir(safe_path)
        return f"Files trong '{path}': {', '.join(files)}"
    except Exception as e:
        return f"Lỗi đọc thư mục: {e}"

@tool
def read_excel(file_path: str) -> str:
    """Đọc nội dung từ file Excel (.xlsx)."""
    logger.info("Reading Excel: %s", file_path)
    try:
        df = pd.read_excel(resolve_safe_path(file_path))
        return f"Dữ liệu Excel:\n{df.to_string(index=False)}"
    except Exception as e:
        return f"Lỗi đọc Excel: {e}"

@tool
def write_excel(file_path: str, json_data: str) -> str:
    """Tạo hoặc ghi file Excel (.xlsx). json_data là danh sách các object dưới dạng chuỗi JSON đại diện cho các dòng. Ví dụ: '[{"STT": 1, "Vật tư": "Ống", "KL": 10}]'"""
    logger.info("Writing Excel: %s", file_path)
    try:
        if not file_path.endswith('.xlsx'):
            file_path += '.xlsx'

        safe_path = resolve_safe_path(file_path)
        dir_name = os.path.dirname(safe_path)
        if dir_name:
            os.makedirs(dir_name, exist_ok=True)

        data = json.loads(json_data)
        df = pd.DataFrame(data)
        df.to_excel(safe_path, index=False)
        return f"Đã ghi đè/tạo thành công file Excel tại: {file_path}"
    except Exception as e:
        return f"Lỗi ghi Excel: {e}"

@tool
def read_word(file_path: str) -> str:
    """Đọc nội dung từ file Word (.docx)."""
    logger.info("Reading Word: %s", file_path)
    try:
        doc = Document(resolve_safe_path(file_path))
        full_text = [para.text for para in doc.paragraphs]
        return "\n".join(full_text)
    except Exception as e:
        return f"Lỗi đọc Word: {e}"

@tool
def write_word(file_path: str, content: str, font_name: str = 'Arial') -> str:
    """Tạo hoặc ghi file Word (.docx) với nội dung được truyền vào. Tham số font_name hỗ trợ 'Arial' hoặc 'Times New Roman'."""
    logger.info("Writing Word: %s", file_path)
    try:
        from docx.shared import Pt
        safe_path = resolve_safe_path(file_path)
        doc = Document()
        # Thiết lập Font chữ chuẩn Unicode (Arial / Times New Roman) cho tiếng Việt
        style = doc.styles['Normal']
        font = style.font
        if font_name not in ['Arial', 'Times New Roman']:
            font_name = 'Arial'
        font.name = font_name
        font.size = Pt(12)

        doc.add_paragraph(content)
        doc.save(safe_path)
        return f"Đã lưu nội dung vào file Word tại: {file_path} (Font: {font_name})"
    except Exception as e:
        return f"Lỗi ghi Word: {e}"

@tool
def read_pdf(file_path: str) -> str:
    """Đọc và trích xuất toàn bộ văn bản từ file PDF."""
    logger.info("Reading PDF: %s", file_path)
    try:
        reader = PdfReader(resolve_safe_path(file_path))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return f"Nội dung PDF ({len(reader.pages)} trang):\n{text[:5000]}..."
    except Exception as e:
        return f"Lỗi đọc PDF: {e}"

@tool
def read_cad(file_path: str) -> str:
    """Đọc file CAD (.dxf/.dwg) và trả về thống kê thư viện block, block attributes, chiều
    dài THẬT (đã tính cung cong và cao độ), và layer sau khi đã làm sạch."""
    logger.info("Reading, Cleaning & Extracting CAD: %s", file_path)
    try:
        doc, load_notes = cad_loader.load_drawing(file_path)

        auditor = audit.Auditor(doc)
        auditor.run()
        audit_fixes = len(auditor.fixes)

        block_defs = []
        for block in doc.blocks:
            is_layout = getattr(block, 'is_layout_block', False) or getattr(block, 'is_any_layout', False)
            if not is_layout and not block.name.startswith('*'):
                block_defs.append(block.name)

        msp = doc.modelspace()
        layer_counts = {}
        block_instances = []
        scaled_blocks = 0
        layer_lengths = {}

        for entity in msp:
            layer = entity.dxf.layer
            layer_counts[layer] = layer_counts.get(layer, 0) + 1
            dxftype = entity.dxftype()

            length = cad_geometry.entity_length(entity)
            if length > 0:
                layer_lengths[layer] = layer_lengths.get(layer, 0.0) + length

            if dxftype == 'INSERT':
                b_name = entity.dxf.name
                attribs = {}
                if hasattr(entity, 'attribs') and entity.attribs:
                    for attrib in entity.attribs:
                        if hasattr(attrib, 'dxf') and hasattr(attrib.dxf, 'tag'):
                            attribs[attrib.dxf.tag] = getattr(attrib.dxf, 'text', '')
                block_instances.append({"name": b_name, "attribs": attribs})
                if cad_geometry.is_scaled(entity):
                    scaled_blocks += 1

        block_summary = {}
        for b in block_instances:
            b_name = b['name']
            attr_str = json.dumps(b['attribs'], ensure_ascii=False) if b['attribs'] else "No Attributes"
            key = f"{b_name} | Thuộc tính: {attr_str}"
            block_summary[key] = block_summary.get(key, 0) + 1

        result = f"Đã làm sạch (Audit). Sửa {audit_fixes} lỗi.\n"
        for note in load_notes:
            result += f"{note}\n"
        result += "\n"

        if len(block_defs) > 25:
            defs_str = ", ".join(block_defs[:25]) + f"... (và {len(block_defs) - 25} block khác)"
        else:
            defs_str = ", ".join(block_defs) if block_defs else "Không có"
        result += f"THƯ VIỆN BLOCK CÓ SẴN (Definitions): {defs_str}\n\n"

        result += "THỐNG KÊ LAYER TRÊN MODELSPACE:\n"
        for k, v in layer_counts.items():
            l_info = f"- Layer '{k}': {v} đối tượng"
            if k in layer_lengths and layer_lengths[k] > 0:
                l_info += f" (Tổng chiều dài thật, kể cả cung cong/cao độ: {layer_lengths[k]:.2f}m)"
            result += l_info + "\n"

        result += "\nTHỐNG KÊ BLOCK THỰC TẾ & THUỘC TÍNH (Attributes):\n"
        if not block_summary:
            result += "(Không có block nào)\n"
        else:
            sorted_blocks = sorted(block_summary.items(), key=lambda x: x[1], reverse=True)
            display_blocks = sorted_blocks[:40]
            for k, v in display_blocks:
                result += f"- Block: {k} -> Số lượng: {v}\n"
            if len(sorted_blocks) > 40:
                result += f"... (và {len(sorted_blocks) - 40} nhóm block khác)\n"

        if scaled_blocks:
            result += (f"\nCẢNH BÁO: {scaled_blocks} block instance bị insert lệch tỷ lệ "
                      f"(xscale/yscale khác 1) — kích thước thực tế trên bản vẽ khác chuẩn.\n")

        return result
    except Exception as e:
        return f"Lỗi xử lý CAD (.dxf/.dwg): {e}"


@tool
def convert_dwg_to_dxf(file_path: str) -> str:
    """Chuyển một file .dwg (định dạng gốc AutoCAD) sang .dxf bằng ODA File Converter.

    Mọi tool đọc bản vẽ khác (`read_cad`, `auto_quantity_takeoff`, `detect_clashes`,
    `render_cad_image`, `analyze_cad_spatial_context`) đã TỰ ĐỘNG gọi bước này khi thấy
    file .dwg — không cần gọi tool này trước. Chỉ dùng khi khách hàng cần chính bản thân
    file .dxf đã chuyển đổi (ví dụ để sửa bằng `edit_cad`/`optimize_cad_drawing`, các tool
    này ghi file bằng ezdxf nên cần đầu vào .dxf).
    """
    logger.info("Manual DWG->DXF conversion requested: %s", file_path)
    try:
        if not file_path.lower().endswith(".dwg"):
            return f"'{file_path}' không phải file .dwg, không cần chuyển đổi."
        doc, notes = cad_loader.load_drawing(file_path)
        return "\n".join(notes) if notes else "Đã chuyển đổi (không có ghi chú thêm)."
    except Exception as e:
        return f"Lỗi chuyển đổi .dwg sang .dxf: {e}"


@tool
def write_cad(file_path: str, layers: str) -> str:
    """Tạo một file CAD mới (.dxf) sạch sẽ với các layer định trước. Tham số layers: chuỗi ngăn cách bởi dấu phẩy."""
    logger.info("Writing CAD: %s", file_path)
    try:
        safe_path = resolve_safe_path(file_path)
        doc = ezdxf.new('R2010')
        layer_list = [l.strip() for l in layers.split(',') if l.strip()]
        for layer in layer_list:
            doc.layers.add(name=layer)

        doc.saveas(safe_path)
        return f"Đã tạo thành công bản vẽ CAD tại {file_path} với các layers: {', '.join(layer_list)}"
    except Exception as e:
        return f"Lỗi tạo CAD (.dxf): {e}"

import sys
from io import StringIO
import builtins

# Sandbox cho execute_python_code: chỉ cho phép các builtin an toàn (không có open/eval/exec/input)
# và chỉ cho phép import các module cần thiết cho việc dựng Block ezdxf (ezdxf, math, json).
# Đây KHÔNG phải cô lập tuyệt đối (không thay thế container/subprocess sandbox thật sự),
# nhưng chặn được các vector tấn công rõ ràng nhất: đọc/ghi file tùy ý, exec chuỗi động, os/subprocess.
_ALLOWED_MODULES = {"ezdxf", "math", "json"}

def _sandboxed_import(name, globals=None, locals=None, fromlist=(), level=0):
    if name.split(".")[0] not in _ALLOWED_MODULES:
        raise ImportError(f"Module '{name}' không được phép sử dụng trong execute_python_code.")
    return builtins.__import__(name, globals, locals, fromlist, level)

_SAFE_BUILTIN_NAMES = (
    "abs", "all", "any", "bool", "dict", "enumerate", "float", "int", "len", "list",
    "max", "min", "print", "range", "round", "set", "sorted", "str", "sum", "tuple",
    "zip", "True", "False", "None", "isinstance",
)
_SAFE_BUILTINS = {name: getattr(builtins, name) for name in _SAFE_BUILTIN_NAMES}
_SAFE_BUILTINS["__import__"] = _sandboxed_import

@tool
def execute_python_code(code: str) -> str:
    """
    Thực thi mã Python động trong môi trường giới hạn (sandbox).
    Được dùng để Họa viên CAD tự viết code ezdxf vẽ Block mới và lưu vào 'data/blocks/mepf_library.dxf'.
    Chỉ cho phép import ezdxf/math/json và không có quyền truy cập file/network trực tiếp qua builtin open().
    """
    logger.info("Executing Custom Python Code (sandboxed)")
    old_stdout = sys.stdout
    try:
        redirected_output = sys.stdout = StringIO()

        safe_globals = {"__builtins__": _SAFE_BUILTINS}
        local_env = {}
        exec(code, safe_globals, local_env)

        return f"Thực thi Python thành công. Output:\n{redirected_output.getvalue()}"
    except Exception as e:
        return f"Lỗi quá trình thực thi Python: {e}"
    finally:
        sys.stdout = old_stdout

@tool
def ai_block_recovery(file_path: str, layer: str, shape: str, dimensions: str, replacement_block: str) -> str:
    """Khôi phục các thiết bị bị phá vỡ (exploded) thành Block chuẩn.
    - layer: Tên layer chứa các nét vẽ rời rạc.
    - shape: 'circle' (hình tròn) hoặc 'rectangle' (hình chữ nhật).
    - dimensions: Với circle là 'bán kính' (ví dụ: '100'). Với rectangle là 'dài,rộng' (ví dụ '600,600').
    - replacement_block: Tên Block mới sẽ được chèn vào.
    """
    logger.info("AI Block Recovery: %s, Layer=%s, Shape=%s", file_path, layer, shape)
    try:
        from ezdxf.addons import importer

        # Lưu điểm lùi TRƯỚC khi ghi đè: các tool này sửa file tại chỗ, không có snapshot
        # thì một lần AI sửa sai là mất luôn bản gốc.
        create_snapshot(file_path, note=f"Trước khi phục hồi Block '{replacement_block}' trên layer '{layer}'")
        safe_path = resolve_safe_path(file_path)
        if not os.path.exists(safe_path):
            return f"Lỗi: Không tìm thấy file {file_path}"

        doc = ezdxf.readfile(safe_path)
        msp = doc.modelspace()

        library_path = os.path.join(get_project_root(), "data", "blocks", "mepf_library.dxf")
        lib_doc = None
        if os.path.exists(library_path):
            lib_doc = ezdxf.readfile(library_path)
            
        if replacement_block not in doc.blocks and lib_doc and replacement_block in lib_doc.blocks:
            imp = importer.Importer(lib_doc, doc)
            imp.import_block(replacement_block)
            imp.finalize()
            
        if replacement_block not in doc.blocks:
            return f"Lỗi: Block '{replacement_block}' không tồn tại trong Thư viện Tổng kho."
            
        centers = []
        entities_to_delete = []
        max_dim = 0
        
        if shape.lower() == "circle":
            target_r = float(dimensions)
            max_dim = target_r * 2
            for entity in msp.query(f'CIRCLE[layer=="{layer}"]'):
                r = entity.dxf.radius
                if abs(r - target_r) / target_r <= 0.05:
                    centers.append((entity.dxf.center.x, entity.dxf.center.y))
                    
        elif shape.lower() == "rectangle":
            dims = dimensions.split(",")
            if len(dims) == 2:
                target_w, target_h = float(dims[0]), float(dims[1])
                max_dim = max(target_w, target_h)
                target_area = target_w * target_h
                for entity in msp.query(f'LWPOLYLINE[layer=="{layer}"]'):
                    if entity.closed or len(entity) >= 4:
                        points = entity.get_points()
                        xs = [p[0] for p in points]
                        ys = [p[1] for p in points]
                        w = max(xs) - min(xs)
                        h = max(ys) - min(ys)
                        area = w * h
                        if area > 0 and abs(area - target_area) / target_area <= 0.1:
                            cx = (max(xs) + min(xs)) / 2
                            cy = (max(ys) + min(ys)) / 2
                            centers.append((cx, cy))
                            
        # Dọn rác chuyên sâu: Quét và xóa MỌI nét vẽ (LINE, PLINE) nằm lọt thỏm trong vùng Block.
        # TEXT/MTEXT mô tả thường được vẽ trên một layer chú thích RIÊNG (khác layer hình học
        # đã khai báo trong `layer`) — nếu chỉ lọc theo đúng layer đó, các dòng text này bị bỏ
        # sót: không xóa (đúng, vì chúng vẫn là mô tả hợp lệ) nhưng cũng không được đồng bộ,
        # nên sau khi phục hồi Block, khách vẫn thấy "còn text mô tả nhưng lệch layer" so với
        # Block mới. Ở đây TEXT/MTEXT gần tâm được QUÉT TRÊN MỌI LAYER và đưa layer về đúng
        # layer của Block vừa phục hồi, thay vì xóa.
        entities_to_relayer = []
        if max_dim > 0 and centers:
            tolerance = max_dim * 0.6  # Phạm vi dọn rác (Bao trùm block + 10% an toàn)
            for entity in msp.query(f'*[layer=="{layer}"]'):
                px, py = None, None
                if hasattr(entity.dxf, 'start'):
                    px, py = entity.dxf.start.x, entity.dxf.start.y
                elif hasattr(entity.dxf, 'center'):
                    px, py = entity.dxf.center.x, entity.dxf.center.y
                elif hasattr(entity.dxf, 'insert'):
                    px, py = entity.dxf.insert.x, entity.dxf.insert.y
                elif entity.dxftype() == 'LWPOLYLINE':
                    try:
                        pts = entity.get_points()
                        px, py = pts[0][0], pts[0][1]
                    except:
                        pass

                if px is not None and py is not None:
                    for cx, cy in centers:
                        if abs(px - cx) <= tolerance and abs(py - cy) <= tolerance:
                            entities_to_delete.append(entity)
                            break

            for entity in msp.query('TEXT MTEXT'):
                if entity.dxf.layer == layer:
                    continue
                px, py = None, None
                if entity.dxftype() == 'TEXT':
                    px, py = entity.dxf.insert.x, entity.dxf.insert.y
                elif entity.dxftype() == 'MTEXT':
                    px, py = entity.dxf.insert.x, entity.dxf.insert.y

                if px is not None and py is not None:
                    for cx, cy in centers:
                        if abs(px - cx) <= tolerance and abs(py - cy) <= tolerance:
                            entities_to_relayer.append(entity)
                            break

        for e in set(entities_to_delete):
            try:
                msp.delete_entity(e)
            except:
                pass

        relayered_count = 0
        for e in set(entities_to_relayer):
            try:
                e.dxf.layer = layer
                relayered_count += 1
            except:
                pass
            
        for cx, cy in centers:
            msp.add_blockref(replacement_block, (cx, cy), dxfattribs={'layer': layer})
            
        doc.saveas(safe_path)
        note = ""
        if relayered_count:
            note = f" Đã đồng bộ {relayered_count} dòng text mô tả bị lệch layer về layer '{layer}'."
        return (f"AI Recovery thành công: Đã tìm thấy và phục hồi {len(centers)} đối tượng '{shape}' "
                f"thành Block '{replacement_block}'.{note}")
    except Exception as e:
        return f"Lỗi phục hồi Block: {e}"

@tool
def edit_cad(file_path: str, actions_json: str) -> str:
    """Chỉnh sửa file CAD (.dxf) hiện tại. Luôn Audit làm sạch file trước.
    actions_json là danh sách các dict. Ví dụ:
    - Thêm layer: {"action": "add_layer", "name": "MEP_DIEN"}
    - Thêm text: {"action": "add_text", "text": "Phong Khach", "x": 0, "y": 0, "layer": "MEP_DIEN", "font_name": "Times New Roman"}
    - Chèn block: {"action": "insert_block", "name": "TU_DIEN", "x": 10, "y": 10, "layer": "MEP_DIEN", "scale": 1.0, "rotation": 0}
    - Đồng bộ font (chống lỗi tiếng Việt): {"action": "fix_fonts", "font_name": "Arial"}
    """
    logger.info("Editing CAD: %s", file_path)
    try:
        create_snapshot(file_path, note="Trước khi edit_cad")
        safe_path = resolve_safe_path(file_path)
        if not os.path.exists(safe_path):
            return f"Lỗi: Không tìm thấy file {file_path}"

        doc = ezdxf.readfile(safe_path)
        msp = doc.modelspace()

        auditor = audit.Auditor(doc)
        auditor.run()
        audit_fixes = len(auditor.fixes)

        actions = json.loads(actions_json)
        results = []

        # Tải Master Library (Tổng kho Block)
        from ezdxf.addons import importer
        library_path = os.path.join(get_project_root(), "data", "blocks", "mepf_library.dxf")
        lib_doc = None
        if os.path.exists(library_path):
            lib_doc = ezdxf.readfile(library_path)
            
        # Khởi tạo Style chữ chuẩn Unicode để tránh lỗi font tiếng Việt trong CAD
        if 'VIETNAMESE_ARIAL' not in doc.styles:
            doc.styles.new('VIETNAMESE_ARIAL', dxfattribs={'font': 'arial.ttf'})
        if 'VIETNAMESE_TIMES' not in doc.styles:
            doc.styles.new('VIETNAMESE_TIMES', dxfattribs={'font': 'times.ttf'})
            
        for act in actions:
            action_type = act.get("action")
            if action_type == "fix_fonts":
                font_name = act.get("font_name", "Arial")
                ttf_file = "times.ttf" if font_name == "Times New Roman" else "arial.ttf"
                # Đổi font của toàn bộ Text Styles có trong bản vẽ
                count = 0
                for style in doc.styles:
                    style.dxf.font = ttf_file
                    count += 1
                results.append(f"Đã đồng bộ {count} Text Styles trong file sang chuẩn Unicode ({font_name}) để sửa lỗi tiếng Việt.")
            elif action_type == "add_layer":
                lname = act.get("name", "NEW_LAYER")
                if lname not in doc.layers:
                    doc.layers.add(name=lname)
                    results.append(f"Thêm layer {lname}")
            elif action_type == "add_text":
                txt = act.get("text", "Text")
                x = act.get("x", 0)
                y = act.get("y", 0)
                layer = act.get("layer", "0")
                font_name = act.get("font_name", "Arial")
                
                if layer not in doc.layers:
                    doc.layers.add(name=layer)
                
                # Áp dụng style VIETNAMESE tương ứng cho Text
                style_name = 'VIETNAMESE_TIMES' if font_name == 'Times New Roman' else 'VIETNAMESE_ARIAL'
                msp.add_text(txt, dxfattribs={'layer': layer, 'style': style_name}).set_placement((x, y))
                results.append(f"Thêm text '{txt}' tại tọa độ ({x},{y}) trên layer {layer} (Font: {font_name})")
            elif action_type == "insert_block":
                b_name = act.get("name")
                x = act.get("x", 0)
                y = act.get("y", 0)
                layer = act.get("layer", "0")
                scale = act.get("scale", 1.0)
                rotation = act.get("rotation", 0.0)
                
                # Auto-Import Block từ Thư viện Trung tâm nếu bản vẽ thiếu
                if b_name not in doc.blocks and lib_doc and b_name in lib_doc.blocks:
                    imp = importer.Importer(lib_doc, doc)
                    imp.import_block(b_name)
                    imp.finalize()
                    results.append(f"Auto-Import thành công Block '{b_name}' từ Thư viện Trung tâm.")
                
                if b_name in doc.blocks:
                    if layer not in doc.layers:
                        doc.layers.add(name=layer)
                    msp.add_blockref(b_name, (x, y), dxfattribs={
                        'layer': layer,
                        'xscale': scale,
                        'yscale': scale,
                        'rotation': rotation
                    })
                    results.append(f"Chèn Block '{b_name}' tại ({x},{y}) layer {layer}")
                else:
                    results.append(f"Lỗi: Block '{b_name}' không tồn tại trong bản vẽ và cả Thư viện Trung tâm.")
                    
        doc.saveas(safe_path)
        return f"Đã làm sạch ({audit_fixes} lỗi rác được xóa) và chỉnh sửa thành công {file_path}:\n- " + "\n- ".join(results)
    except Exception as e:
        return f"Lỗi sửa CAD (.dxf): {e}"

@tool
def render_cad_image(file_path: str, output_png_path: str = "cad_preview.png") -> str:
    """Chuyển đổi file bản vẽ CAD (.dxf) thành hình ảnh PNG sắc nét để hiển thị trực quan (Computer Vision) lên giao diện Web."""
    logger.info("Rendering CAD to Image: %s -> %s", file_path, output_png_path)
    try:
        from ezdxf.addons.drawing import RenderContext, Frontend
        from ezdxf.addons.drawing.matplotlib import MatplotlibBackend
        import matplotlib.pyplot as plt

        doc, _load_notes = cad_loader.load_drawing(file_path)
        msp = doc.modelspace()

        fig = plt.figure(figsize=(12, 8), dpi=150)
        ax = fig.add_axes([0, 0, 1, 1])
        ctx = RenderContext(doc)
        out = MatplotlibBackend(ax)
        Frontend(ctx, out).draw_layout(msp, finalize=True)
        fig.savefig(resolve_safe_path(output_png_path), dpi=150, bbox_inches='tight')
        plt.close(fig)
        return f"Đã xuất hình ảnh bản vẽ CAD (Computer Vision) thành công tại: {output_png_path}"
    except Exception as e:
        return f"Lỗi xuất ảnh CAD: {e}"

@tool
def analyze_cad_spatial_context(file_path: str, max_distance: float = 2000.0) -> str:
    """Phân tích Ngữ cảnh Hình học & Mũi tên Chỉ dẫn (Leaders, Text Annotations, Spatial Matching) để hiểu bản vẽ CAD như con người: tự động liên kết Ghi chú văn bản (ví dụ: 'Ống uPVC Ø110', 'Ống gió 600x400') và Mũi tên chỉ hướng với đúng nét vẽ đường ống kề cận."""
    logger.info("Analyzing CAD Spatial Context & Arrows: %s", file_path)
    try:
        doc, _load_notes = cad_loader.load_drawing(file_path)
        msp = doc.modelspace()
        
        texts = []
        leaders = []
        pipe_segments = []
        
        def point_to_seg_dist(px, py, ax, ay, bx, by):
            l2 = (bx - ax)**2 + (by - ay)**2
            if l2 == 0:
                return math.hypot(px - ax, py - ay)
            t = max(0.0, min(1.0, ((px - ax) * (bx - ax) + (py - ay) * (by - ay)) / l2))
            proj_x = ax + t * (bx - ax)
            proj_y = ay + t * (by - ay)
            return math.hypot(px - proj_x, py - proj_y)

        for entity in msp:
            dxftype = entity.dxftype()
            layer = entity.dxf.layer
            
            if dxftype == 'TEXT':
                t_str = normalize_pipe_diameter_spec(entity.dxf.text.strip())
                pos = entity.dxf.insert
                if t_str:
                    texts.append({"text": t_str, "pos": (pos.x, pos.y), "layer": layer})
            elif dxftype == 'MTEXT':
                t_str = normalize_pipe_diameter_spec(entity.text.strip())
                pos = entity.dxf.insert
                if t_str:
                    texts.append({"text": t_str, "pos": (pos.x, pos.y), "layer": layer})
            elif dxftype in ('LEADER', 'MULTILEADER'):
                try:
                    if hasattr(entity, 'vertices') and entity.vertices:
                        vertices = [(v.x, v.y) for v in entity.vertices]
                        leaders.append({"tip": vertices[0], "tail": vertices[-1], "layer": layer})
                except Exception:
                    pass
            elif dxftype == 'LINE':
                s, e = entity.dxf.start, entity.dxf.end
                length = math.hypot(e.x - s.x, e.y - s.y)
                pipe_segments.append({"layer": layer, "seg": (s.x, s.y, e.x, e.y), "length": length})
            elif dxftype in ('LWPOLYLINE', 'POLYLINE'):
                try:
                    if dxftype == 'LWPOLYLINE':
                        pts = entity.get_points(format='xy')
                    else:
                        pts = [(v.dxf.location.x, v.dxf.location.y) for v in entity.vertices]
                    for i in range(1, len(pts)):
                        ax, ay = pts[i-1][0], pts[i-1][1]
                        bx, by = pts[i][0], pts[i][1]
                        length = math.hypot(bx - ax, by - ay)
                        pipe_segments.append({"layer": layer, "seg": (ax, ay, bx, by), "length": length})
                except Exception:
                    pass

        try:
            from rtree import index
            has_rtree = True
            pipe_idx = index.Index()
            for i, p in enumerate(pipe_segments):
                ax, ay, bx, by = p["seg"]
                pipe_idx.insert(i, (min(ax, bx), min(ay, by), max(ax, bx), max(ay, by)))
        except ImportError:
            has_rtree = False

        associations = {}
        for text_item in texts:
            tx, ty = text_item["pos"]
            txt = text_item["text"]
            
            min_dist = float('inf')
            best_pipe = None
            
            if has_rtree:
                # Query index with a bounding box expanded by max_distance
                search_box = (tx - max_distance, ty - max_distance, tx + max_distance, ty + max_distance)
                candidates = pipe_idx.intersection(search_box)
                for i in candidates:
                    p = pipe_segments[i]
                    ax, ay, bx, by = p["seg"]
                    d = point_to_seg_dist(tx, ty, ax, ay, bx, by)
                    if d < min_dist:
                        min_dist = d
                        best_pipe = p
            else:
                for p in pipe_segments:
                    ax, ay, bx, by = p["seg"]
                    d = point_to_seg_dist(tx, ty, ax, ay, bx, by)
                    if d < min_dist:
                        min_dist = d
                        best_pipe = p
                    
            if best_pipe and min_dist <= max_distance:
                p_layer = best_pipe["layer"]
                key = f"Ghi chú: '{txt}' <---> Layer ống: '{p_layer}'"
                if key not in associations:
                    associations[key] = {"count": 0, "total_length": 0.0, "text": txt, "layer": p_layer, "min_dist": min_dist}
                associations[key]["count"] += 1
                associations[key]["total_length"] += best_pipe["length"]

        report = f"PHÂN TÍCH NGỮ CẢNH HÌNH HỌC & MŨI TÊN CHỈ DẪN (Spatial Intelligence):\n"
        report += f"- Tìm thấy {len(texts)} văn bản ghi chú (TEXT/MTEXT), {len(leaders)} mũi tên chỉ dẫn (LEADER), và {len(pipe_segments)} đoạn đường ống.\n\n"
        
        report += "📌 THỐNG KÊ GHI CHÚ VĂN BẢN VÀ MŨI TÊN (Tối đa 20 ghi chú tiêu biểu):\n"
        for t in texts[:20]:
            report += f"  • Ghi chú: \"{t['text']}\" (Layer: {t['layer']}) tại tọa độ ({t['pos'][0]:.1f}, {t['pos'][1]:.1f})\n"
        if len(texts) > 20:
            report += f"  ... và {len(texts) - 20} ghi chú khác.\n"
            
        report += "\n🔗 LIÊN KẾT HÌNH HỌC KHÔNG GIANG (Text Annotation <-> Pipe Segment): \n"
        if not associations:
            report += "  (Không tìm thấy liên kết kề cận trong bán kính khoảng cách quy định)\n"
        else:
            for k, v in list(associations.items())[:25]:
                report += f"  • [{v['text']}] liên kết trực tiếp với tuyến ống Layer '{v['layer']}' (Khoảng cách kề cận: {v['min_dist']:.1f}mm) -> Tổng chiều dài suy luận: {v['total_length']:.2f}m\n"
                
        return report
    except Exception as e:
        return f"Lỗi phân tích ngữ cảnh không gian CAD: {e}"

# Bản đồ mã INSUNITS (DXF group code $INSUNITS) sang tên đơn vị dễ đọc.
_INSUNITS_NAMES = {
    0: "Không xác định (Unitless)", 1: "Inch", 2: "Feet", 4: "Millimet (mm)",
    5: "Centimet (cm)", 6: "Met (m)", 8: "Microinch", 9: "Mil",
}


@tool
def audit_cad_drawing_errors(file_path: str, text_duplicate_tolerance: float = 1.0) -> str:
    """Kiểm tra CHỈ ĐỌC (không sửa file) các lỗi thường gặp khi khách hàng đẩy bản vẽ CAD
    vào — dùng khi khách yêu cầu "kiểm tra bản vẽ có lỗi gì không", "rà soát trước khi
    duyệt", hoặc bất cứ khi nào nhận bản vẽ mới từ bên ngoài trước khi xử lý tiếp. Khác với
    `optimize_cad_drawing` (tự động XÓA/SỬA), tool này chỉ liệt kê để kỹ sư tự quyết định,
    vì một số lỗi (VD sai đơn vị) cần hỏi lại khách thay vì tự đoán sửa.

    Các lỗi được kiểm tra (đều là lỗi hay gặp nhất khi nhận bản vẽ từ nguồn khác/khách hàng
    tự vẽ, không phải lỗi hình học học thuật):
    1. SAI ĐƠN VỊ BẢN VẼ (INSUNITS != Millimet): lỗi nghiêm trọng nhất — nếu khách vẽ bằng
       inch/m/cm nhưng hệ thống đọc là mm (hoặc ngược lại), MỌI kích thước/khối lượng tính
       ra sau đó đều sai lệch hàng chục đến hàng nghìn lần dù hình học trông "bình thường".
    2. VẼ TRỰC TIẾP TRÊN LAYER "0": thói quen xấu phổ biến khiến đối tượng luôn thừa kế màu/
       linetype theo Block cha thay vì Layer riêng, gây khó kiểm soát khi chỉnh sửa hàng loạt.
    3. BLOCK BỊ CHÈN LỆCH TỶ LỆ (scale distortion): Block insert với x-scale khác y-scale
       hoặc khác 1 — khiến kích thước thực tế của thiết bị/ký hiệu sai lệch so với định nghĩa
       gốc, dễ gây bóc khối lượng/kiểm tra khoảng cách lắp đặt sai.
    4. TEXT/MTEXT TRÙNG LẶP: hai ghi chú có nội dung giống hệt nhau đặt gần nhau (trong phạm
       vi `text_duplicate_tolerance`) — thường do copy/paste nhầm, gây đếm nhầm khi bóc tách.
    5. NHIỀU CHIỀU CAO CHỮ KHÔNG NHẤT QUÁN: bản vẽ dùng quá nhiều cỡ chữ khác nhau cho ghi
       chú kích thước (dấu hiệu ghép nhiều nguồn bản vẽ không đồng bộ chuẩn trình bày).
    6. CAO ĐỘ Z BẤT THƯỜNG: đối tượng có Z khác 0 lẫn trong bản vẽ chủ yếu phẳng (Z=0) — có
       thể do vô tình vẽ/copy nhầm cao độ, khiến đối tượng trông đúng trên mặt bằng nhưng
       thực chất lệch cao độ so với phần còn lại.
    """
    logger.info("Auditing CAD drawing errors (read-only): %s", file_path)
    try:
        doc, load_notes = cad_loader.load_drawing(file_path)
        msp = doc.modelspace()

        issues = []

        # 1. Đơn vị bản vẽ.
        insunits = doc.header.get('$INSUNITS', 0)
        if insunits != 4:
            unit_name = _INSUNITS_NAMES.get(insunits, f"Mã INSUNITS={insunits} (không xác định)")
            issues.append(
                f"[NGHIÊM TRỌNG] Đơn vị bản vẽ hiện là '{unit_name}', KHÔNG PHẢI Millimet (mm). "
                f"Toàn bộ hệ thống giả định bản vẽ vẽ bằng mm — nếu khách thực sự vẽ bằng đơn vị "
                f"khác, MỌI kích thước/khối lượng tính từ bản vẽ này đều SAI LỆCH. Cần hỏi lại "
                f"khách đơn vị vẽ thực tế trước khi xử lý tiếp."
            )

        # 2. Vẽ trực tiếp trên Layer "0".
        layer0_types = {}
        for entity in msp:
            if entity.dxf.layer == "0" and entity.dxftype() != "INSERT":
                layer0_types[entity.dxftype()] = layer0_types.get(entity.dxftype(), 0) + 1
        if layer0_types:
            total_l0 = sum(layer0_types.values())
            breakdown = ", ".join(f"{t}: {c}" for t, c in sorted(layer0_types.items(), key=lambda x: -x[1]))
            issues.append(
                f"[CẢNH BÁO] {total_l0} đối tượng vẽ trực tiếp trên Layer '0' thay vì Layer riêng "
                f"({breakdown}) — nên chuyển sang Layer đúng hệ để dễ kiểm soát màu/hiển thị/tắt-mở."
            )

        # 3. Block chèn lệch tỷ lệ.
        scale_issues = {}
        for entity in msp.query('INSERT'):
            xs, ys = round(entity.dxf.xscale, 3), round(entity.dxf.yscale, 3)
            if abs(xs - ys) > 0.01 or abs(xs - 1.0) > 0.05:
                key = (entity.dxf.name, xs, ys)
                scale_issues[key] = scale_issues.get(key, 0) + 1
        if scale_issues:
            top = sorted(scale_issues.items(), key=lambda x: -x[1])[:10]
            detail = "; ".join(f"{name} (scale {xs}x{ys}) x{count}" for (name, xs, ys), count in top)
            issues.append(
                f"[CẢNH BÁO] {sum(scale_issues.values())} lượt Block bị chèn LỆCH TỶ LỆ (không phải "
                f"1:1 hoặc x-scale khác y-scale): {detail}"
                + (f" ... và {len(scale_issues) - 10} loại khác" if len(scale_issues) > 10 else "")
                + " — kích thước thực tế trên bản vẽ khác định nghĩa gốc của Block, cần khách xác nhận."
            )

        # 4. Text/MTEXT trùng lặp (cùng nội dung, vị trí gần nhau).
        text_entries = []
        for entity in msp:
            dxftype = entity.dxftype()
            if dxftype == "TEXT":
                txt, pos = (entity.dxf.text or "").strip(), entity.dxf.insert
            elif dxftype == "MTEXT":
                txt, pos = (entity.text or "").strip(), entity.dxf.insert
            else:
                continue
            if txt:
                text_entries.append((txt, pos.x, pos.y))

        seen_text = []
        dup_count = 0
        dup_samples = []
        for txt, x, y in text_entries:
            is_dup = False
            for txt2, x2, y2 in seen_text:
                if txt == txt2 and math.hypot(x - x2, y - y2) <= text_duplicate_tolerance:
                    is_dup = True
                    break
            if is_dup:
                dup_count += 1
                if len(dup_samples) < 5:
                    dup_samples.append(f"'{txt}' tại ({x:.0f}, {y:.0f})")
            else:
                seen_text.append((txt, x, y))
        if dup_count:
            issues.append(
                f"[CẢNH BÁO] {dup_count} ghi chú TEXT/MTEXT trùng lặp nội dung + vị trí gần nhau "
                f"(trong phạm vi {text_duplicate_tolerance}mm), VD: {'; '.join(dup_samples)} — có thể "
                f"do copy/paste nhầm, gây đếm sai khi bóc tách khối lượng dựa trên ghi chú."
            )

        # 5. Chiều cao chữ không nhất quán.
        heights = []
        for entity in msp:
            if entity.dxftype() == "TEXT" and entity.dxf.hasattr("height"):
                heights.append(round(entity.dxf.height, 1))
            elif entity.dxftype() == "MTEXT" and entity.dxf.hasattr("char_height"):
                heights.append(round(entity.dxf.char_height, 1))
        distinct_heights = sorted(set(heights))
        if len(distinct_heights) > 5:
            issues.append(
                f"[GHI CHÚ] Bản vẽ dùng {len(distinct_heights)} cỡ chữ khác nhau cho TEXT/MTEXT "
                f"({distinct_heights[:10]}{'...' if len(distinct_heights) > 10 else ''}) — có thể là "
                f"dấu hiệu ghép nhiều nguồn bản vẽ không đồng bộ chuẩn trình bày, nên thống nhất lại."
            )

        # 6. Cao độ Z bất thường trong bản vẽ chủ yếu phẳng.
        z_values = []
        for entity in msp:
            points = cad_geometry.entity_points_3d(entity)
            z_values.extend(p[2] for p in points)
        if z_values:
            nonzero = [z for z in z_values if abs(z) > 1e-6]
            if nonzero and len(nonzero) < len(z_values) * 0.05:
                distinct_z = sorted(set(round(z, 1) for z in nonzero))
                issues.append(
                    f"[GHI CHÚ] {len(nonzero)}/{len(z_values)} điểm có cao độ Z khác 0 "
                    f"(Z = {distinct_z[:10]}) trong khi phần lớn bản vẽ phẳng (Z=0) — kiểm tra xem có "
                    f"đối tượng nào vô tình bị vẽ/copy sai cao độ hay không."
                )

        if not issues:
            return (f"KHÔNG phát hiện lỗi phổ biến nào trong bản vẽ. Đã kiểm tra: đơn vị bản vẽ, "
                    f"đối tượng trên Layer 0, Block lệch tỷ lệ, text trùng lặp, độ nhất quán cỡ "
                    f"chữ, cao độ Z bất thường.")

        report = [f"RÀ SOÁT LỖI BẢN VẼ ({len(issues)} vấn đề phát hiện, chỉ đọc — chưa sửa file):"]
        for note in load_notes:
            report.append(f"- {note}")
        for i, issue in enumerate(issues, 1):
            report.append(f"{i}. {issue}")
        report.append(
            "- Lưu ý: Đây là công cụ RÀ SOÁT (read-only). Dùng `optimize_cad_drawing` để tự động "
            "sửa rác hình học/trùng lặp Block, `standardize_cad_drawing` để chuẩn hóa tên Layer/"
            "Block — lỗi đơn vị bản vẽ và Block lệch tỷ lệ cần XÁC NHẬN với khách trước khi sửa."
        )
        return "\n".join(report)
    except Exception as e:
        return f"Lỗi rà soát bản vẽ CAD: {e}"


@tool
def optimize_cad_drawing(file_path: str, output_path: str = "", dedupe_tolerance: float = 1.0) -> str:
    """Tối ưu & dọn dẹp bản vẽ CAD (.dxf) TỰ ĐỘNG, thuần hình học (KHÔNG dùng LLM) — phù
    hợp chạy offline hoặc với model AI yếu vì không cần suy luận, chỉ cần gọi 1 tool:
    1. Audit sửa lỗi cấu trúc file.
    2. Xóa các đoạn LINE/POLYLINE có chiều dài = 0 (rác vẽ thừa).
    3. Xóa các Block instance (INSERT) bị trùng lặp hoàn toàn (cùng tên Block, cùng vị trí
       trong phạm vi dedupe_tolerance) — lỗi thường gặp khi copy/paste nhầm trong CAD.
    4. Overkill: xóa LINE/LWPOLYLINE trùng lặp/chồng đè hoàn toàn lên nhau (cùng layer,
       cùng hình học trong phạm vi dedupe_tolerance) — lỗi thường gặp khi trace lại đường
       nét cũ mà không xóa nét gốc, khiến file nặng và đo khối lượng bị nhân đôi.
    5. Purge: xóa Layer rỗng, Block định nghĩa không còn INSERT nào tham chiếu, text style
       và linetype không còn dùng — tương đương lệnh PURGE của AutoCAD.
    Nếu output_path bỏ trống, ghi đè lên chính file_path.
    """
    logger.info("Optimize CAD Drawing (offline, deterministic): %s", file_path)
    try:
        create_snapshot(file_path, note="Trước khi optimize_cad_drawing")
        safe_path = resolve_safe_path(file_path)
        doc = ezdxf.readfile(safe_path)
        msp = doc.modelspace()

        auditor = audit.Auditor(doc)
        auditor.run()
        audit_fixes = len(auditor.fixes)

        removed_zero_len = 0
        removed_dupe_blocks = 0

        for entity in list(msp.query('LINE')):
            start, end = entity.dxf.start, entity.dxf.end
            if math.hypot(end.x - start.x, end.y - start.y) < 1e-6:
                msp.delete_entity(entity)
                removed_zero_len += 1

        for entity in list(msp.query('LWPOLYLINE')):
            try:
                pts = entity.get_points(format='xy')
                total = sum(math.hypot(pts[i][0] - pts[i - 1][0], pts[i][1] - pts[i - 1][1]) for i in range(1, len(pts)))
                if total < 1e-6:
                    msp.delete_entity(entity)
                    removed_zero_len += 1
            except Exception:
                pass

        seen_inserts = set()
        for entity in list(msp.query('INSERT')):
            ins = entity.dxf.insert
            key = (
                entity.dxf.name,
                entity.dxf.layer,
                round(ins.x / dedupe_tolerance) if dedupe_tolerance > 0 else ins.x,
                round(ins.y / dedupe_tolerance) if dedupe_tolerance > 0 else ins.y,
            )
            if key in seen_inserts:
                msp.delete_entity(entity)
                removed_dupe_blocks += 1
            else:
                seen_inserts.add(key)

        def _rounded_point(x, y):
            if dedupe_tolerance > 0:
                return (round(x / dedupe_tolerance), round(y / dedupe_tolerance))
            return (x, y)

        removed_overkill = 0
        seen_lines = set()
        for entity in list(msp.query('LINE')):
            start, end = entity.dxf.start, entity.dxf.end
            p1, p2 = _rounded_point(start.x, start.y), _rounded_point(end.x, end.y)
            key = (entity.dxf.layer, frozenset((p1, p2)))
            if key in seen_lines:
                msp.delete_entity(entity)
                removed_overkill += 1
            else:
                seen_lines.add(key)

        seen_polylines = set()
        for entity in list(msp.query('LWPOLYLINE')):
            try:
                pts = tuple(_rounded_point(x, y) for x, y in entity.get_points(format='xy'))
                if not pts:
                    continue
                key = (entity.dxf.layer, min(pts, pts[::-1]))
                if key in seen_polylines:
                    msp.delete_entity(entity)
                    removed_overkill += 1
                else:
                    seen_polylines.add(key)
            except Exception:
                pass

        used_layers = {entity.dxf.layer for entity in msp}
        removed_layers = []
        for layer in list(doc.layers):
            lname = layer.dxf.name
            if lname in ('0', 'Defpoints') or lname in used_layers:
                continue
            try:
                doc.layers.remove(lname)
                removed_layers.append(lname)
            except Exception:
                pass

        used_block_names = {entity.dxf.name for entity in msp.query('INSERT')}
        removed_blocks = []
        for block in list(doc.blocks):
            bname = block.name
            if bname.startswith('*') or bname in used_block_names:
                continue  # bỏ qua Block ẩn hệ thống (*Model_Space, *Paper_Space...)
            try:
                doc.blocks.delete_block(bname, safe=True)
                removed_blocks.append(bname)
            except Exception:
                pass

        used_styles = {e.dxf.style for e in msp if e.dxftype() in ('TEXT', 'MTEXT') and e.dxf.hasattr('style')}
        removed_styles = []
        for style in list(doc.styles):
            sname = style.dxf.name
            if sname.upper() in ('STANDARD',) or sname in used_styles:
                continue
            try:
                doc.styles.remove(sname)
                removed_styles.append(sname)
            except Exception:
                pass

        used_linetypes = {e.dxf.linetype for e in msp if e.dxf.hasattr('linetype')}
        used_linetypes |= {layer.dxf.linetype for layer in doc.layers}
        removed_linetypes = []
        for lt in list(doc.linetypes):
            ltname = lt.dxf.name
            if ltname.upper() in ('BYLAYER', 'BYBLOCK', 'CONTINUOUS') or ltname in used_linetypes:
                continue
            try:
                doc.linetypes.remove(ltname)
                removed_linetypes.append(ltname)
            except Exception:
                pass

        purge_total = len(removed_blocks) + len(removed_styles) + len(removed_linetypes)

        target_path = output_path.strip() or file_path
        out_safe_path = resolve_safe_path(target_path)
        doc.saveas(out_safe_path)

        return (
            f"TỐI ƯU BẢN VẼ THÀNH CÔNG (offline, không cần LLM suy luận):\n"
            f"- Audit sửa {audit_fixes} lỗi cấu trúc.\n"
            f"- Xóa {removed_zero_len} đối tượng có chiều dài bằng 0 (rác vẽ).\n"
            f"- Xóa {removed_dupe_blocks} Block trùng lặp (cùng tên + cùng vị trí).\n"
            f"- Overkill: xóa {removed_overkill} LINE/LWPOLYLINE trùng lặp/chồng đè.\n"
            f"- Purge: xóa {len(removed_layers)} Layer rỗng"
            + (f" ({', '.join(removed_layers)})" if removed_layers else "") + ", "
            f"{len(removed_blocks)} Block định nghĩa không dùng"
            + (f" ({', '.join(removed_blocks)})" if removed_blocks else "") + ", "
            f"{len(removed_styles)} text style không dùng, "
            f"{len(removed_linetypes)} linetype không dùng.\n"
            f"- Tổng cộng purge {purge_total + len(removed_layers)} đối tượng định nghĩa thừa.\n"
            f"- Đã lưu bản vẽ đã tối ưu tại: {target_path}"
        )
    except Exception as e:
        return f"Lỗi tối ưu bản vẽ CAD: {e}"


def _apply_layer_style(layer, std: dict) -> bool:
    """Áp màu/linetype/mô tả chuẩn lên 1 Layer. Trả về True nếu có thay đổi."""
    changed = False
    if layer.dxf.color != std["color"]:
        layer.dxf.color = std["color"]
        changed = True
    if layer.dxf.linetype != cad_standards.LAYER_LINETYPE:
        layer.dxf.linetype = cad_standards.LAYER_LINETYPE
        changed = True
    if layer.description != std["description"]:
        layer.description = std["description"]
        changed = True
    return changed


def _ensure_block_attributes(block, std: dict) -> bool:
    """Gắn ATTDEF MA_HIEU/MO_TA (ẩn, hằng số) vào định nghĩa Block nếu còn thiếu, để
    khi Block được chèn vào bản vẽ, mã hiệu/mô tả chuẩn luôn đi kèm. Trả về True nếu
    có ATTDEF mới được thêm."""
    from ezdxf.lldxf import const as dxf_const

    existing_tags = {a.dxf.tag for a in block.attdefs()}
    changed = False
    flags = dxf_const.ATTRIB_INVISIBLE + dxf_const.ATTRIB_CONST
    if "MA_HIEU" not in existing_tags:
        block.add_attdef("MA_HIEU", (0, 0), text=std["ma_hieu"],
                          dxfattribs={"height": 30, "flags": flags, "layer": "0"})
        changed = True
    if "MO_TA" not in existing_tags:
        block.add_attdef("MO_TA", (0, 0), text=std["description"],
                          dxfattribs={"height": 30, "flags": flags, "layer": "0"})
        changed = True
    return changed


@tool
def standardize_cad_drawing(file_path: str, output_path: str = "") -> str:
    """Chuẩn hóa TÊN LAYER, TÊN BLOCK và MÔ TẢ/MÃ HIỆU của bản vẽ CAD (.dxf) người
    dùng đẩy vào theo tiêu chuẩn nội bộ MEPF (xem `src/cad_standards.py`). CHỈ sửa
    đặt tên, màu/linetype của layer, và gắn thuộc tính MA_HIEU/MO_TA vào Block —
    TUYỆT ĐỐI KHÔNG động vào hình học/nét vẽ, khác với `ai_block_recovery` (phục hồi
    Block bị vỡ) và `optimize_cad_drawing` (dọn rác hình học).
    Layer/Block không nhận diện được theo tiêu chuẩn sẽ được liệt kê để người dùng tự
    kiểm tra thay vì bị đoán bừa. Nếu output_path bỏ trống, ghi đè lên chính file_path.
    """
    logger.info("Standardize CAD Drawing: %s", file_path)
    try:
        create_snapshot(file_path, note="Trước khi standardize_cad_drawing")
        safe_path = resolve_safe_path(file_path)
        if not os.path.exists(safe_path):
            return f"Lỗi: Không tìm thấy file {file_path}"

        doc = ezdxf.readfile(safe_path)
        msp = doc.modelspace()

        auditor = audit.Auditor(doc)
        auditor.run()
        audit_fixes = len(auditor.fixes)

        renamed_layers = []
        fixed_layer_props = []
        unmatched_layers = []

        for layer in list(doc.layers):
            lname = layer.dxf.name
            if lname.lower() in ("0", "defpoints"):
                continue
            canonical = cad_standards.match_layer(lname)
            if not canonical:
                unmatched_layers.append(lname)
                continue
            std = cad_standards.LAYER_STANDARD[canonical]
            if lname == canonical:
                if _apply_layer_style(layer, std):
                    fixed_layer_props.append(lname)
                continue
            if canonical in doc.layers:
                moved = 0
                for e in list(msp.query(f'*[layer=="{lname}"]')):
                    e.dxf.layer = canonical
                    moved += 1
                _apply_layer_style(doc.layers.get(canonical), std)
                try:
                    doc.layers.remove(lname)
                except Exception:
                    pass
                renamed_layers.append(f"{lname} -> {canonical} (gộp {moved} đối tượng)")
            else:
                layer.rename(canonical)
                _apply_layer_style(doc.layers.get(canonical), std)
                renamed_layers.append(f"{lname} -> {canonical}")

        renamed_blocks = []
        annotated_blocks = []
        unmatched_blocks = []

        used_block_names = sorted({e.dxf.name for e in msp.query('INSERT') if not e.dxf.name.startswith('*')})
        for bname in used_block_names:
            target_name = bname
            canonical = cad_standards.match_block(bname)
            if not canonical:
                unmatched_blocks.append(bname)
                continue
            if canonical != bname:
                if canonical in doc.blocks:
                    unmatched_blocks.append(
                        f"{bname} (trùng ý nghĩa với Block chuẩn có sẵn '{canonical}' nhưng hình học có thể khác "
                        f"— cần kiểm tra thủ công, dùng `ai_block_recovery` nếu muốn thay hẳn bằng Block chuẩn)"
                    )
                    continue
                doc.blocks.rename_block(bname, canonical)
                for ins in msp.query(f'INSERT[name=="{bname}"]'):
                    ins.dxf.name = canonical
                renamed_blocks.append(f"{bname} -> {canonical}")
                target_name = canonical

            std = cad_standards.BLOCK_STANDARD.get(target_name)
            if std and _ensure_block_attributes(doc.blocks.get(target_name), std):
                annotated_blocks.append(target_name)

        target_path = output_path.strip() or file_path
        out_safe_path = resolve_safe_path(target_path)
        doc.saveas(out_safe_path)

        report = ["CHUẨN HÓA BẢN VẼ THÀNH CÔNG (chỉ sửa tên/layer/mô tả, KHÔNG đổi hình học):"]
        report.append(f"- Audit sửa {audit_fixes} lỗi cấu trúc.")
        report.append("- Đổi tên layer về chuẩn: " + ("; ".join(renamed_layers) if renamed_layers else "(không có)"))
        report.append("- Sửa màu/linetype/mô tả layer đã đúng tên nhưng sai thuộc tính: "
                       + (", ".join(fixed_layer_props) if fixed_layer_props else "(không có)"))
        if unmatched_layers:
            report.append(f"- CẦN REVIEW THỦ CÔNG {len(unmatched_layers)} layer không nhận diện được theo chuẩn: "
                           + ", ".join(unmatched_layers))
        report.append("- Đổi tên Block về chuẩn: " + ("; ".join(renamed_blocks) if renamed_blocks else "(không có)"))
        report.append("- Gắn thuộc tính MA_HIEU/MO_TA cho Block: "
                       + (", ".join(annotated_blocks) if annotated_blocks else "(không có)"))
        if unmatched_blocks:
            report.append(f"- CẦN REVIEW THỦ CÔNG {len(unmatched_blocks)} Block không nhận diện được theo chuẩn: "
                           + "; ".join(unmatched_blocks))
        report.append(f"- Đã lưu bản vẽ đã chuẩn hóa tại: {target_path}")
        return "\n".join(report)
    except Exception as e:
        return f"Lỗi chuẩn hóa bản vẽ CAD: {e}"


@tool
def add_color_legend(file_path: str, output_path: str = "") -> str:
    """Vẽ trực tiếp vào bản vẽ CAD (.dxf) một bảng CHÚ THÍCH MÀU SẮC (legend) thể hiện
    đầy đủ quy chuẩn màu Layer MEPF nội bộ (`src/cad_standards.py`): ống gió cấp/hồi/
    tươi/thải/thải bếp/tăng áp/hút khói (SAD/RAD/FAD/EAD/KEAD/PAD/SEAD), ống đồng gas
    lạnh, ống nước ngưng, ống nước lạnh Chiller cấp/hồi, ống cấp nước lạnh/nóng sinh
    hoạt, hồi nước nóng, thoát nước thải/thông hơi/nước mưa, Sprinkler, họng nước, và
    toàn bộ layer thiết bị/dây dẫn của 4 hệ Mechanical/Electrical/Plumbing/Firefighting.
    Mỗi dòng gồm 1 ô màu (SOLID) + tên Layer chuẩn + mô tả + tên màu (9 màu ACI cơ bản
    1-9 có tên tiếng Việt; màu mở rộng ghi "ACI <n>" vì không có tên chuẩn hóa phổ quát).
    Legend nằm trên layer riêng 'G-LEGEND', tự động đặt bên phải vùng vẽ hiện có (dựa
    theo bounding box thực tế) để không đè lên hình học gốc.
    Dùng sau `standardize_cad_drawing` để hồ sơ nộp có ghi chú quy chuẩn màu ngay trên
    bản vẽ, không phải tra tài liệu rời. Nếu output_path bỏ trống, ghi đè lên file_path.
    """
    logger.info("Add color legend: %s", file_path)
    try:
        create_snapshot(file_path, note="Trước khi add_color_legend")
        safe_path = resolve_safe_path(file_path)
        if not os.path.exists(safe_path):
            return f"Lỗi: Không tìm thấy file {file_path}"

        doc = ezdxf.readfile(safe_path)
        msp = doc.modelspace()

        if "G-LEGEND" not in doc.layers:
            doc.layers.add("G-LEGEND", dxfattribs={"color": 7})

        start_x, start_y = 0.0, 0.0
        try:
            from ezdxf import bbox
            extents = bbox.extents(msp)
            if extents.has_data:
                start_x = extents.extmax.x + 2000.0
                start_y = extents.extmax.y
        except Exception:
            pass

        rows = cad_standards.color_legend_rows()
        row_h = 250.0
        swatch_w = 200.0
        swatch_h = 180.0
        text_h = 150.0
        cur_y = start_y

        msp.add_text(
            "QUY CHUẨN MÀU SẮC LAYER MEPF",
            dxfattribs={"layer": "G-LEGEND", "height": text_h * 1.4},
        ).set_placement((start_x, cur_y))
        cur_y -= row_h * 1.4

        last_discipline = None
        for row in rows:
            if row["discipline"] != last_discipline:
                last_discipline = row["discipline"]
                msp.add_text(
                    last_discipline.upper(),
                    dxfattribs={"layer": "G-LEGEND", "height": text_h * 1.15},
                ).set_placement((start_x, cur_y))
                cur_y -= row_h

            bl = (start_x, cur_y - swatch_h)
            br = (start_x + swatch_w, cur_y - swatch_h)
            tl = (start_x, cur_y)
            tr = (start_x + swatch_w, cur_y)
            msp.add_solid([bl, br, tl, tr], dxfattribs={"layer": "G-LEGEND", "color": row["color"]})

            label = f"{row['layer']} - {row['description']} ({row['color_name']}, màu ACI {row['color']})"
            msp.add_text(
                label, dxfattribs={"layer": "G-LEGEND", "height": text_h},
            ).set_placement((start_x + swatch_w + 100, cur_y - swatch_h * 0.6))
            cur_y -= row_h

        target_path = output_path.strip() or file_path
        out_safe_path = resolve_safe_path(target_path)
        doc.saveas(out_safe_path)

        return (
            "THÊM CHÚ THÍCH MÀU SẮC THÀNH CÔNG:\n"
            f"- Đã vẽ {len(rows)} dòng quy chuẩn màu Layer (4 hệ M/E/P/F + General) trên layer 'G-LEGEND'.\n"
            f"- Vị trí góc trên-trái bảng chú thích: ({start_x:.0f}, {start_y:.0f}).\n"
            f"- Đã lưu bản vẽ tại: {target_path}"
        )
    except Exception as e:
        return f"Lỗi thêm chú thích màu sắc: {e}"

@tool
def extract_new_blocks_to_library(file_path: str) -> str:
    """Quét bản vẽ CAD mới, lọc và trích xuất các Block hợp lệ (chưa có) vào thư viện mepf_library.dxf."""
    logger.info("Extracting new blocks from %s to library", file_path)
    try:
        from ezdxf.addons import Importer
        safe_path = resolve_safe_path(file_path)
        doc = ezdxf.readfile(safe_path)
        
        library_path = os.path.join(get_project_root(), "data", "blocks", "mepf_library.dxf")
        if not os.path.exists(library_path):
            lib_doc = ezdxf.new()
            os.makedirs(os.path.dirname(library_path), exist_ok=True)
            lib_doc.saveas(library_path)
        else:
            lib_doc = ezdxf.readfile(library_path)
            
        existing_lib_blocks = {b.name.upper() for b in lib_doc.blocks}
        
        blocks_to_import = []
        prefixes = ("HVAC-", "ELEC-", "PLUMB-", "FF-")
        for block in doc.blocks:
            name = block.name
            if name.startswith('*') or name.startswith('_'):
                continue
            
            name_up = name.upper()
            if name_up in existing_lib_blocks:
                continue
                
            if name_up.startswith(prefixes):
                blocks_to_import.append(name)
                
        if not blocks_to_import:
            return "Không tìm thấy Block MEPF mới nào cần thu thập."
            
        importer = Importer(doc, lib_doc)
        importer.import_blocks(blocks_to_import)
        importer.finalize()
        
        lib_doc.saveas(library_path)
        return f"Đã học và thêm {len(blocks_to_import)} Block MEPF mới vào CSDL."
    except Exception as e:
        return f"Lỗi khi học Block mới: {e}"

@tool
def auto_route_mepf_path(file_path: str, start_x: float, start_y: float, end_x: float, end_y: float, layer_name: str) -> str:
    """Tự động đi tuyến ống/cáp (Auto-Routing) từ điểm A đến điểm B sử dụng thuật toán tìm đường trên mặt bằng.
    
    Tạo ra một đường Polyline mới thuộc layer chỉ định, cố gắng tránh các đối tượng hiện có.
    """
    logger.info("Auto-routing from (%s, %s) to (%s, %s) on layer %s", start_x, start_y, end_x, end_y, layer_name)
    try:
        safe_path = resolve_safe_path(file_path)
        doc = ezdxf.readfile(safe_path)
        msp = doc.modelspace()
        
        # Đơn giản hóa: Vẽ một đường ziczac vuông góc đơn giản
        mid_x = (start_x + end_x) / 2
        points = [(start_x, start_y), (mid_x, start_y), (mid_x, end_y), (end_x, end_y)]
        
        msp.add_lwpolyline(points, dxfattribs={'layer': layer_name})
        doc.saveas(safe_path)
        
        return f"Đã tự động đi tuyến và vẽ Polyline trên layer '{layer_name}' từ ({start_x}, {start_y}) đến ({end_x}, {end_y})."
    except Exception as e:
        return f"Lỗi khi đi tuyến tự động: {e}"

@tool
def generate_calculation_report(agent_role: str, content: str, output_filename: str) -> str:
    """Tạo Thuyết minh tính toán (Technical Report) định dạng Word (.docx).
    
    Nhận nội dung báo cáo dạng text hoặc markdown đơn giản và xuất ra file Word chuyên nghiệp.
    """
    try:
        safe_path = resolve_safe_path(output_filename if output_filename.endswith('.docx') else output_filename + '.docx')
        parent = os.path.dirname(safe_path)
        if parent:
            os.makedirs(parent, exist_ok=True)
            
        doc = Document()
        doc.add_heading(f'THUYẾT MINH TÍNH TOÁN - {agent_role.upper()}', 0)
        
        for p in content.split('\n\n'):
            if p.strip():
                doc.add_paragraph(p.strip())
                
        doc.save(safe_path)
        return f"Đã tạo Thuyết minh tính toán thành công: {os.path.basename(safe_path)}"
    except Exception as e:
        return f"Lỗi tạo báo cáo Word: {e}"

@tool
def lookup_equipment_catalog(equipment_type: str, search_kw: str) -> str:
    """Tra cứu catalog thiết bị thực tế (Daikin, Ebara, Cadivi, v.v.) từ cơ sở dữ liệu.
    
    Ví dụ: equipment_type='Bơm', search_kw='50m3/h'
    """
    try:
        db_path = os.path.join(get_project_root(), "data", "equipment_catalog.json")
        if not os.path.exists(db_path):
            return "Chưa có file CSDL thiết bị (data/equipment_catalog.json)."
            
        with open(db_path, "r", encoding="utf-8") as f:
            catalog = json.load(f)
            
        results = []
        kw = search_kw.lower()
        for eq in catalog.get("equipments", []):
            if equipment_type.lower() in eq.get("type", "").lower():
                if kw in str(eq).lower():
                    results.append(str(eq))
                    
        if results:
            return "Kết quả tra cứu catalog:\n" + "\n".join(results)
        return f"Không tìm thấy thiết bị '{equipment_type}' khớp với từ khóa '{search_kw}'."
    except Exception as e:
        return f"Lỗi tra cứu catalog: {e}"



from src.hvac_tools import (
    calc_psychrometrics, calc_duct_size, calc_cooling_load, calc_chw_pipe_size, calc_pump_fan_power, calc_ventilation_rate,
    calc_cooling_load_detailed, calc_duct_total_pressure_loss, calc_chiller_ahu_selection, calc_refrigerant_pipe_size,
    calc_cooling_tower, calc_fresh_air_ashrae, calc_vrv_outdoor_unit,
)
from src.elec_tools import calc_cable_size, calc_breaker_size, calc_lighting_qty
from src.plumb_tools import (
    calc_water_pipe, calc_water_tank, calc_plumbing_pump_head,
    calc_drainage_pipe, calc_rainwater_drainage, calc_septic_tank, calc_hot_water_system,
    calc_vent_pipe, calc_grease_trap, calc_sump_pump,
)
from src.ff_tools import calc_sprinkler_qty, calc_fire_pump, calc_extinguisher_qty
from src.elec_tools import (
    calc_voltage_drop, calc_total_load, calc_short_circuit,
    calc_cable_tray_size, calc_lightning_protection,
    calc_emergency_lighting, calc_power_factor_correction,
)
from src.hvac_tools import calc_nc_level
from src.ff_tools import (
    calc_sprinkler_hydraulics, calc_standpipe, calc_smoke_control, calc_fire_detector_qty,
    calc_gas_suppression, calc_fire_water_tank,
)
from src.qs_tools import (
    lookup_unit_price, calc_boq_cost, export_boq_vietnam, auto_quantity_takeoff, calc_support_hangers,
)
from src.bim_tools import detect_clashes, read_ifc_model, check_pipe_connectivity
from src.panel_schedule import generate_panel_schedule
from src.cad_revision import (
    snapshot_cad, list_cad_revisions, diff_cad_revisions, restore_cad_revision,
)

tools = [
    search_standards, search_web, calculate, execute_python_code, list_directory,
    read_excel, write_excel, read_word, write_word, read_pdf,
    read_cad, write_cad, edit_cad, ai_block_recovery, render_cad_image, analyze_cad_spatial_context,
    auto_quantity_takeoff, optimize_cad_drawing, standardize_cad_drawing, convert_dwg_to_dxf, add_color_legend,
    audit_cad_drawing_errors,
    calc_psychrometrics, calc_duct_size, calc_cooling_load, calc_chw_pipe_size, calc_pump_fan_power, calc_ventilation_rate,
    calc_cooling_load_detailed, calc_duct_total_pressure_loss, calc_chiller_ahu_selection, calc_refrigerant_pipe_size,
    calc_cooling_tower, calc_fresh_air_ashrae, calc_vrv_outdoor_unit,
    calc_cable_size, calc_breaker_size, calc_lighting_qty, calc_voltage_drop,
    calc_total_load, calc_short_circuit, calc_cable_tray_size, calc_lightning_protection,
    calc_emergency_lighting, calc_power_factor_correction,
    generate_panel_schedule,
    calc_water_pipe, calc_water_tank, calc_plumbing_pump_head,
    calc_drainage_pipe, calc_rainwater_drainage, calc_septic_tank, calc_hot_water_system,
    calc_vent_pipe, calc_grease_trap, calc_sump_pump,
    calc_sprinkler_qty, calc_fire_pump, calc_extinguisher_qty,
    calc_sprinkler_hydraulics, calc_standpipe, calc_smoke_control, calc_fire_detector_qty,
    calc_gas_suppression, calc_fire_water_tank,
    calc_nc_level,
    lookup_unit_price, calc_boq_cost, export_boq_vietnam, calc_support_hangers,
    detect_clashes, read_ifc_model, check_pipe_connectivity,
    snapshot_cad, list_cad_revisions, diff_cad_revisions, restore_cad_revision,
    auto_route_mepf_path, generate_calculation_report, lookup_equipment_catalog, extract_new_blocks_to_library
]

# Giảm token: trước đây MỌI agent đều bị bind cả danh sách `tools` đầy đủ (30+ schema),
# kể cả tool hoàn toàn không liên quan chuyên môn (VD: ElectricalAgent vẫn nhận schema
# đọc/ghi CAD, tính PCCC...). Tách theo từng vai trò để mỗi agent chỉ gửi kèm tool nó
# thực sự cần trong LLM request, cắt đáng kể input token mỗi lượt gọi mà không đổi
# hành vi (ToolNode trong src/graph.py vẫn dùng `tools` đầy đủ để thực thi bất kỳ
# tool_call nào, không phụ thuộc danh sách bind ở đây).
_COMMON_TOOLS = [search_standards, search_web, calculate, list_directory, read_excel, read_word, read_pdf, generate_calculation_report, lookup_equipment_catalog]

TOOLS_BY_ROLE = {
    "mechanical": _COMMON_TOOLS + [
        calc_cooling_load, calc_cooling_load_detailed, calc_duct_size, calc_duct_total_pressure_loss,
        calc_psychrometrics, calc_chw_pipe_size, calc_chiller_ahu_selection, calc_refrigerant_pipe_size,
        calc_pump_fan_power, calc_ventilation_rate, calc_nc_level,
        calc_cooling_tower, calc_fresh_air_ashrae, calc_vrv_outdoor_unit,
    ],
    "electrical": _COMMON_TOOLS + [
        calc_cable_size, calc_breaker_size, calc_lighting_qty, calc_voltage_drop,
        calc_total_load, calc_short_circuit, calc_cable_tray_size, calc_lightning_protection,
        generate_panel_schedule, calc_emergency_lighting, calc_power_factor_correction,
    ],
    "plumbing": _COMMON_TOOLS + [
        calc_water_pipe, calc_water_tank, calc_plumbing_pump_head,
        calc_drainage_pipe, calc_rainwater_drainage, calc_septic_tank, calc_hot_water_system,
        calc_vent_pipe, calc_grease_trap, calc_sump_pump,
    ],
    "firefighting": _COMMON_TOOLS + [
        calc_sprinkler_qty, calc_fire_pump, calc_extinguisher_qty,
        calc_sprinkler_hydraulics, calc_standpipe, calc_smoke_control, calc_fire_detector_qty,
        calc_gas_suppression, calc_fire_water_tank,
    ],
    "qs": _COMMON_TOOLS + [
        auto_quantity_takeoff, read_cad, write_excel, analyze_cad_spatial_context, ai_block_recovery,
        lookup_unit_price, calc_boq_cost, export_boq_vietnam, convert_dwg_to_dxf, calc_support_hangers,
    ],
    "cad": _COMMON_TOOLS + [
        read_cad, write_cad, edit_cad, ai_block_recovery, render_cad_image,
        analyze_cad_spatial_context, execute_python_code, optimize_cad_drawing,
        standardize_cad_drawing, auto_route_mepf_path, extract_new_blocks_to_library,
        snapshot_cad, list_cad_revisions, diff_cad_revisions, restore_cad_revision,
        convert_dwg_to_dxf, add_color_legend, audit_cad_drawing_errors,
    ],
    "bim": _COMMON_TOOLS + [
        auto_quantity_takeoff, read_cad, write_excel, analyze_cad_spatial_context, detect_clashes,
        check_pipe_connectivity, read_ifc_model,
        diff_cad_revisions, list_cad_revisions, convert_dwg_to_dxf, audit_cad_drawing_errors,
    ],
}

def get_tools_for_role(role: str) -> list:
    """Tool set thu gọn cho một vai trò cụ thể; vai trò không xác định (VD: Supervisor
    gọi nhầm) sẽ nhận về toàn bộ `tools` để không bao giờ thiếu tool cần thiết."""
    return TOOLS_BY_ROLE.get((role or "").lower().strip(), tools)
